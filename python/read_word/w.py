from docx import Document
import re
from docx.enum.style import WD_STYLE_TYPE






class GetBiaoZhun():

    def __init__(self,docFile):
        self.doc=Document(docFile)

    def getScene(self):
        for h in self.doc.paragraphs:
            if h.style.name == "Heading 2":
                if re.search("^4.\w+场景[1-9]{0,3}\w+",h.text):
                    print("scene: %s\n" %(h.text))

    def getRequest(self):
        for r in self.doc.paragraphs:
            if r.style.name == "Heading 2":
                if re.search("^4.[0-9]{1,3}.\w+要求",r.text):
                    print("%s" %(r.text))

    def getContext(self):
        singe = self.doc.paragraphs[0]
        print(singe.text)
        # for r in self.doc.paragraphs:
        #     if r.style.name == "Heading 2":
        #         if re.search("^4.[0-9]{1,3}.\w+要求",r.text):
        #             print("%s" %(r))
if __name__=="__main__":
    # 场景名称
    run = GetBiaoZhun("biaozhun.docx")

    run.getContext()